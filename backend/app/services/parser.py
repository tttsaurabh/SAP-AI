import os
import io
import csv
import shutil
import pandas as pd
from typing import Dict, Any, List, Tuple
from loguru import logger
from app.core.config import settings

# Import document parsing libraries with fallback wrappers
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

class DocumentParser:
    @staticmethod
    def parse_file(file_path: str, filename: str) -> List[Dict[str, Any]]:
        """
        Parses a file and returns a list of pages/segments.
        Each item is a dictionary: {"page": int, "text": str, "metadata": dict}
        """
        ext = os.path.splitext(filename)[1].lower()
        logger.info(f"Parsing file {filename} with extension {ext}")
        
        if ext == ".pdf":
            if settings.PDF_PARSER_ENGINE == "unlimited_ocr":
                return DocumentParser._parse_pdf_unlimited_ocr(file_path)
            return DocumentParser._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentParser._parse_docx(file_path)
        elif ext in [".xlsx", ".xls", ".csv"]:
            return DocumentParser._parse_spreadsheet(file_path, ext)
        elif ext in [".txt", ".md", ".markdown"]:
            return DocumentParser._parse_text(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return DocumentParser._parse_image_ocr(file_path)
        else:
            # Fallback to plain text read
            try:
                return DocumentParser._parse_text(file_path)
            except Exception as e:
                logger.error(f"Failed to parse file {filename}: {str(e)}")
                raise ValueError(f"Unsupported file format and fallback text parsing failed: {ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> List[Dict[str, Any]]:
        pages = []
        
        # 1. Try PyMuPDF first (fast and robust)
        if fitz:
            try:
                doc = fitz.open(file_path)
                for page_idx, page in enumerate(doc):
                    text = page.get_text("text") or ""
                    # Check if we should try OCR for scanned/empty pages
                    if len(text.strip()) < 50 and pytesseract and Image:
                        # Try OCR, but wrap in try-except so OCR unavailability doesn't fail the whole document
                        try:
                            pix = page.get_pixmap()
                            img_data = pix.tobytes("png")
                            img = Image.open(io.BytesIO(img_data))
                            ocr_text = pytesseract.image_to_string(img)
                            if len(ocr_text.strip()) > len(text.strip()):
                                text = ocr_text
                        except Exception as ocr_e:
                            logger.debug(f"OCR failed for page {page_idx + 1}: {str(ocr_e)}")
                    
                    pages.append({
                        "page": page_idx + 1,
                        "text": text,
                        "metadata": {
                            "source_type": "pdf",
                            "engine": "pymupdf",
                            "headings": DocumentParser._extract_headings_from_text(text)
                        }
                    })
                doc.close()
                if any(len(p["text"].strip()) > 0 for p in pages):
                    return pages
            except Exception as e:
                logger.warning(f"PyMuPDF parsing failed, falling back: {str(e)}")
        
        # 2. Try pdfplumber for table preservation
        if pdfplumber:
            try:
                pages_plumber = []
                with pdfplumber.open(file_path) as pdf:
                    for idx, page in enumerate(pdf.pages):
                        text = page.extract_text() or ""
                        # Extract tables
                        tables = page.extract_tables()
                        table_str = ""
                        if tables:
                            for table in tables:
                                # Convert table grid to markdown table
                                table_str += "\n\n"
                                for r_idx, row in enumerate(table):
                                    row_filtered = [str(cell or "").replace("\n", " ").strip() for cell in row]
                                    table_str += "| " + " | ".join(row_filtered) + " |\n"
                                    if r_idx == 0:
                                        table_str += "| " + " | ".join(["---"] * len(row_filtered)) + " |\n"
                                table_str += "\n"
                        
                        full_text = text + table_str
                        pages_plumber.append({
                            "page": idx + 1,
                            "text": full_text,
                            "metadata": {
                                "source_type": "pdf",
                                "engine": "pdfplumber",
                                "has_tables": len(tables) > 0
                            }
                        })
                return pages_plumber
            except Exception as e:
                logger.warning(f"pdfplumber parsing failed: {str(e)}")
        
        # Fallback empty list or basic parsing if libraries failed
        if not pages:
            pages.append({
                "page": 1,
                "text": f"[Parsing failed or libraries not available for {os.path.basename(file_path)}]",
                "metadata": {"error": "No parser engines completed"}
            })
        return pages

    @staticmethod
    def _parse_docx(file_path: str) -> List[Dict[str, Any]]:
        if not docx:
            # Fallback to plain text extraction of DOCX (using zip if possible, or warning)
            return [{
                "page": 1,
                "text": "[python-docx library not installed. Cannot parse DOCX]",
                "metadata": {"source_type": "docx", "error": "docx package missing"}
            }]
        
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Paragraphs
            for p in doc.paragraphs:
                if p.text:
                    full_text.append(p.text)
            
            # Tables
            for table in doc.tables:
                full_text.append("\n")
                for r_idx, row in enumerate(table.rows):
                    row_cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    # Simple deduplication because Word tables can duplicate cells in columns
                    full_text.append("| " + " | ".join(row_cells) + " |")
                    if r_idx == 0:
                        full_text.append("| " + " | ".join(["---"] * len(row_cells)) + " |")
                full_text.append("\n")

            text_content = "\n".join(full_text)
            return [{
                "page": 1,
                "text": text_content,
                "metadata": {
                    "source_type": "docx",
                    "headings": DocumentParser._extract_headings_from_text(text_content)
                }
            }]
        except Exception as e:
            logger.error(f"Error parsing Word document: {str(e)}")
            raise e

    @staticmethod
    def _parse_spreadsheet(file_path: str, ext: str) -> List[Dict[str, Any]]:
        try:
            pages = []
            if ext == ".csv":
                df = pd.read_csv(file_path)
                markdown_table = df.to_markdown(index=False)
                pages.append({
                    "page": 1,
                    "text": markdown_table,
                    "metadata": {"source_type": "csv", "rows": len(df)}
                })
            else:
                # Excel file - process sheet by sheet
                xl = pd.ExcelFile(file_path)
                for idx, sheet_name in enumerate(xl.sheet_names):
                    df = xl.parse(sheet_name)
                    markdown_table = df.to_markdown(index=False)
                    pages.append({
                        "page": idx + 1,
                        "text": f"Sheet: {sheet_name}\n\n{markdown_table}",
                        "metadata": {"source_type": "excel", "sheet_name": sheet_name, "rows": len(df)}
                    })
            return pages
        except Exception as e:
            logger.error(f"Error parsing spreadsheet: {str(e)}")
            raise e

    @staticmethod
    def _parse_text(file_path: str) -> List[Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return [{
                "page": 1,
                "text": content,
                "metadata": {
                    "source_type": "text",
                    "headings": DocumentParser._extract_headings_from_text(content)
                }
            }]
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            raise e

    @staticmethod
    def _parse_image_ocr(file_path: str) -> List[Dict[str, Any]]:
        if not pytesseract or not Image:
            return [{
                "page": 1,
                "text": "[OCR required libraries (Tesseract / Pillow) not available]",
                "metadata": {"source_type": "image", "error": "tesseract packages missing"}
            }]
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return [{
                "page": 1,
                "text": text,
                "metadata": {"source_type": "image", "ocr": True}
            }]
        except Exception as e:
            logger.error(f"Error performing OCR: {str(e)}")
            raise e

    @staticmethod
    def _extract_headings_from_text(text: str) -> List[str]:
        headings = []
        for line in text.split("\n"):
            line = line.strip()
            # Catch Markdown headings
            if line.startswith("#"):
                headings.append(line.lstrip("#").strip())
            # Catch uppercase lines that look like structural SAP titles
            elif line.isupper() and 3 < len(line) < 100:
                headings.append(line)
        return headings[:15] # Cap headings at 15 per page for metadata size

    @staticmethod
    def _parse_pdf_unlimited_ocr(file_path: str) -> List[Dict[str, Any]]:
        """
        Parses a PDF using Baidu's Unlimited-OCR transformer model.
        1. Render PDF pages as images in a temporary folder.
        2. Load tokenizer and AutoModel from 'baidu/Unlimited-OCR'.
        3. Run inference via model.infer for each page image.
        4. Return pages with parsed markdown text.
        """
        import tempfile
        from PIL import Image
        
        logger.info(f"Initiating Baidu Unlimited-OCR parsing for {file_path}")
        
        # Ensure PyMuPDF is available to render pages to images
        if not fitz:
            raise ImportError("PyMuPDF is required to render PDF pages as images for Unlimited-OCR. Please install pymupdf.")
            
        try:
            import torch
            from transformers import AutoModel, AutoTokenizer
        except ImportError:
            raise ImportError("PyTorch and Transformers are required for Baidu Unlimited-OCR. Please install torch and transformers.")
            
        doc = fitz.open(file_path)
        pages_count = len(doc)
        
        # Render pages to temporary PNG images
        temp_dir = tempfile.mkdtemp()
        image_paths = []
        try:
            for page_idx in range(pages_count):
                page = doc[page_idx]
                pix = page.get_pixmap(dpi=150) # render with reasonable resolution
                img_path = os.path.join(temp_dir, f"page_{page_idx}.png")
                pix.save(img_path)
                image_paths.append(img_path)
                
            logger.info(f"Rendered {pages_count} pages as images for OCR in: {temp_dir}")
            
            # Load Baidu Unlimited-OCR model
            model_name = 'baidu/Unlimited-OCR'
            logger.info(f"Loading {model_name} model and tokenizer...")
            tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
            
            # Use GPU if available, else fallback to CPU
            device = "cuda" if torch.cuda.is_available() else "cpu"
            dtype = torch.bfloat16 if torch.cuda.is_available() else torch.float32
            
            model = AutoModel.from_pretrained(
                model_name, 
                trust_remote_code=True, 
                use_safetensors=True, 
                torch_dtype=dtype
            ).eval().to(device)
            
            logger.info(f"Baidu Unlimited-OCR model loaded successfully on device: {device}")
            
            pages_data = []
            
            # Page-by-page inference for detailed per-page tracking
            for page_idx, img_path in enumerate(image_paths):
                logger.info(f"OCR processing page {page_idx + 1}/{pages_count}...")
                
                # Single page inference config: base mode
                out_dir = os.path.join(temp_dir, f"output_{page_idx}")
                os.makedirs(out_dir, exist_ok=True)
                
                # The model's custom infer method writes prediction to a text file
                model.infer(
                    tokenizer,
                    prompt='<image>document parsing.',
                    image_file=img_path,
                    output_path=out_dir,
                    base_size=1024,
                    image_size=1024,
                    crop_mode=False,
                    max_length=32768,
                    save_results=True
                )
                
                # Read the generated result from output path
                page_text = ""
                if os.path.exists(out_dir):
                    out_files = os.listdir(out_dir)
                    for f_name in out_files:
                        if f_name.endswith(".txt"):
                            with open(os.path.join(out_dir, f_name), "r", encoding="utf-8") as f_ref:
                                page_text = f_ref.read()
                                break
                            
                # Fallback prediction description if file wasn't written
                if not page_text:
                    logger.warning(f"No text output file generated by Unlimited-OCR for page {page_idx + 1}")
                    page_text = f"[Baidu Unlimited-OCR could not extract text from page {page_idx + 1}]"
                    
                pages_data.append({
                    "page": page_idx + 1,
                    "text": page_text,
                    "metadata": {
                        "source_type": "pdf",
                        "engine": "baidu_unlimited_ocr",
                        "headings": DocumentParser._extract_headings_from_text(page_text)
                    }
                })
                
            return pages_data
            
        except Exception as e:
            logger.error(f"Baidu Unlimited-OCR parsing failed: {str(e)}")
            raise e
            
        finally:
            # Clean up temp folder and images
            doc.close()
            try:
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary OCR images from: {temp_dir}")
            except Exception as clean_err:
                logger.warning(f"Failed to clean up temp OCR folder: {str(clean_err)}")

